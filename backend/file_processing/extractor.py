# file_processing/extractor.py
import os
from PyPDF2 import PdfReader
from docx import Document
import markdown # For converting MD to HTML then to text, or just read directly

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts text content from a PDF file.
    """
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n" # Add newline between pages
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        # Consider re-raising or returning an empty string/error indicator
        raise ValueError(f"Could not extract text from PDF: {e}")
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text content from a DOCX file.
    """
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}")
    return text.strip()

def extract_text_from_md(file_path: str) -> str:
    """
    Extracts text content from a Markdown file.
    For simplicity, this reads the raw Markdown.
    One could convert MD to HTML then strip tags for a cleaner text version if needed.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        # Optional: Convert Markdown to plain text more robustly
        # html = markdown.markdown(text)
        # from bs4 import BeautifulSoup
        # soup = BeautifulSoup(html, 'html.parser')
        # return soup.get_text()
    except Exception as e:
        print(f"Error reading MD {file_path}: {e}")
        raise ValueError(f"Could not extract text from Markdown: {e}")
    return text.strip()

def extract_text(file_path: str, original_filename: str) -> str:
    """
    Detects file type from original_filename and extracts text accordingly.
    """
    _, extension = os.path.splitext(original_filename.lower())
    
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension == '.docx':
        return extract_text_from_docx(file_path)
    elif extension == '.md':
        return extract_text_from_md(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}. Supported types are .pdf, .docx, .md.")

if __name__ == '__main__':
    # Create dummy files for testing
    # Ensure the 'storage/uploads' directory exists relative to the project root for this test
    # This test assumes it's run from the project root or paths are adjusted.
    # For module-level testing, you might place test files within the module or use absolute paths.

    # Simplified test - assumes you run this script from rag_service/file_processing directory
    # and have test_files/ in rag_service/file_processing
    test_files_dir = "test_files" # Create this directory in file_processing/
    os.makedirs(test_files_dir, exist_ok=True)
    
    # PDF Test
    pdf_test_path = os.path.join(test_files_dir, "test.pdf")
    # Note: PyPDF2 cannot create PDFs, so you'd need a sample PDF.
    # For now, we'll just show how to call it.
    # print(f"To test PDF, place a test.pdf in {test_files_dir}")
    # if os.path.exists(pdf_test_path):
    #     print("\n--- PDF Test ---")
    #     try:
    #         pdf_text = extract_text(pdf_test_path, "test.pdf")
    #         print(f"Extracted from PDF (first 100 chars): {pdf_text[:100]}")
    #     except ValueError as e:
    #         print(e)
    # else:
    #     print(f"Skipping PDF test: {pdf_test_path} not found.")


    # DOCX Test
    docx_test_path = os.path.join(test_files_dir, "test.docx")
    try:
        doc = Document()
        doc.add_paragraph("This is a test document for DOCX extraction.")
        doc.add_paragraph("It has multiple paragraphs.")
        doc.save(docx_test_path)
        print("\n--- DOCX Test ---")
        docx_text = extract_text(docx_test_path, "test.docx")
        print(f"Extracted from DOCX: {docx_text}")
    except Exception as e: # Catching general exception for docx library issues if it's not installed
        print(f"Could not run DOCX test (ensure python-docx is installed and no permission issues): {e}")


    # MD Test
    md_test_path = os.path.join(test_files_dir, "test.md")
    with open(md_test_path, 'w', encoding='utf-8') as f:
        f.write("# Markdown Test\n\nThis is a test markdown file.\n\n- Item 1\n- Item 2")
    print("\n--- MD Test ---")
    md_text = extract_text(md_test_path, "test.md")
    print(f"Extracted from MD: {md_text}")

    # Unsupported file type test
    print("\n--- Unsupported File Test ---")
    try:
        extract_text("dummy.txt", "dummy.txt")
    except ValueError as e:
        print(f"Caught expected error: {e}")
